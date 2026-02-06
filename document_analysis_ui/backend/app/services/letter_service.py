"""Service for generating Word document letters for students."""

import os
from pathlib import Path
from docx import Document
from docx.shared import Pt
import structlog

logger = structlog.get_logger(__name__)

class LetterService:
    """Service to generate professional letters in Word format."""

    def generate_admission_letter(
        self, 
        student_data: dict, 
        output_path: Path, 
        override_name: str | None = None,
        override_date: str | None = None
    ) -> Path | None:
        """
        Generate an admission eligibility letter based on a Word template.
        
        Args:
            student_data: Data extracted by the orchestrator.
            output_path: Path where the .docx file should be saved.
            override_name: Manual name override from user.
            override_date: Manual date override from user.
            
        Returns:
            Path to the generated file, or None if generation failed.
        """
        try:
            # Locate template in project root
            template_path = Path(__file__).parent.parent.parent.parent.parent / "conditional letter_v1.docx"
            
            if not template_path.exists():
                logger.error("template_not_found", path=str(template_path))
                return None
                
            doc = Document(template_path)
            
            # 1. Determine the name
            if override_name:
                formatted_name = override_name.strip().title()
            else:
                passport = student_data.get("passport_details", {})
                first_name = (passport.get("first_name") or "").strip()
                last_name = (passport.get("last_name") or "").strip()
                raw_full_name = f"{first_name} {last_name}".strip()
                
                if raw_full_name:
                    formatted_name = raw_full_name.title()
                else:
                    formatted_name = "Student"
            
            # 2. Determine the date
            if override_date:
                formatted_date = override_date.strip()
            else:
                # Requirement: January 6th, 2026
                from datetime import datetime
                now = datetime.now()
                month = now.strftime("%B")
                day = now.day
                year = now.year
                
                # Add ordinal suffix
                if 11 <= day <= 13:
                    suffix = "th"
                else:
                    suffix = {1: "st", 2: "nd", 3: "rd"}.get(day % 10, "th")
                
                formatted_date = f"{month} {day}{suffix}, {year}"
            
            logger.info("generating_letter_content", name=formatted_name, date=formatted_date)

            # Replace placeholders in paragraphs
            for p in doc.paragraphs:
                self._replace_placeholders(p, formatted_name, formatted_date)
            
            # Replace in tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            self._replace_placeholders(p, formatted_name, formatted_date)

            # Save the document
            doc.save(str(output_path))
            logger.info("letter_generated", path=str(output_path), student=formatted_name)
            return output_path
            
        except Exception as e:
            logger.error("failed_to_generate_letter", error=str(e))
            return None

    def _replace_placeholders(self, p, name: str, date_str: str):
        """Helper to replace {{name}} and {{date}} in a paragraph with split-run awareness."""
        placeholders = {
            '{{name}}': name,
            '{{date}}': date_str
        }
        
        has_split_placeholder = False
        any_placeholder_found = False
        
        for placeholder, replacement in placeholders.items():
            if placeholder in p.text:
                any_placeholder_found = True
                replaced_in_runs = False
                for run in p.runs:
                    if placeholder in run.text:
                        run.text = run.text.replace(placeholder, replacement)
                        replaced_in_runs = True
                
                if not replaced_in_runs:
                    has_split_placeholder = True

        # If any placeholder was split across runs, we must do a paragraph-level replacement
        # to ensure it's replaced, even if it's destructive to formatting.
        if has_split_placeholder:
            text = p.text
            for placeholder, replacement in placeholders.items():
                text = text.replace(placeholder, replacement)
            p.text = text

# Global instance
letter_service = LetterService()
